import os
from datetime import datetime
from pathlib import Path

from docx.shared import Inches
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

BASE_DIR = Path(__file__).resolve().parent


class CreateDocument:
    def __init__(self, filename, image_width):
        self.my_doc = docx.Document()
        self.filename = filename
        self.image_width = image_width

    def set_data(self, initial_data, table):
        for i in range(0, len(initial_data), 2):
            row_cells_label = table.add_row().cells
            row_cells_image = table.add_row().cells

            if len(initial_data) == 1:
                row_cells_label[0].merge(row_cells_label[1])
                row_cells_image[0].merge(row_cells_image[1])

            row_cells_label[0].text = f"{initial_data[i][0]}"
            paragraph = row_cells_image[0].paragraphs[0]
            paragraph.alignment = 1
            run = paragraph.add_run()
            run.add_picture(initial_data[i][1], width=self.image_width)

            if i + 1 < len(initial_data):
                row_cells_label[1].text = f"{initial_data[i + 1][0]}"
                paragraph = row_cells_image[1].paragraphs[0]
                paragraph.alignment = 1
                run = paragraph.add_run()
                run.add_picture(initial_data[i + 1][1], width=self.image_width)
        # self.table.add_row()

    def set_file_data(self, initial_data):
        file_names = sorted(initial_data.keys())
        for file_name_without_extension in file_names:
            table = self.my_doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'

            file_data = initial_data[file_name_without_extension]
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].text = file_name_without_extension
            p = row_cells[0].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            self.set_data(file_data, table)

            if file_name_without_extension != file_names[-1]:
                self.my_doc.add_page_break()

        self.my_doc.save(str(BASE_DIR / self.filename))


if __name__ == '__main__':
    current_time = datetime.now().time()

    args = (
        ('My 1', r"/home/root366/Dump/GITHUB/Images-to-Docs/abc"),
        ('My 2', r"/home/root366/Dump/GITHUB/Images-to-Docs/vcc"),
        ('My 3', r"/home/root366/Dump/GITHUB/Images-to-Docs/mvp"),
    )
    filename = f'testing_{current_time.hour}_{current_time.minute}_{current_time.second}.docx'

    data = {}
    for folder_name, folder_path in args:
        files = os.listdir(folder_path)
        for file in files:
            file_name_without_ext = file.split('.')[0]
            data[file_name_without_ext] = data.get(file_name_without_ext, []) + [
                [folder_name, os.path.join(folder_path, file)]]

    width = Inches(2.8)

    current_time = datetime.now().time()
    create_docs = CreateDocument(filename=filename, image_width=width)
    create_docs.set_file_data(data)
